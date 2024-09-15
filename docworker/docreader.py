from docx import Document


def read_docx_in_paragraphs(file_path):
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    return paragraphs


def read_table_in_paragraphs(file_path):
    doc = Document(file_path)
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cells.append(cell.text)
    return cells


